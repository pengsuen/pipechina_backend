from __future__ import annotations

import hashlib
import io
import json
from datetime import UTC, datetime, time, timedelta
from uuid import UUID
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from docx import Document
from sqlalchemy import and_, func, or_, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.handover.domain.models import AudioRecord, HandoverSummaryVersion
from app.modules.inspection.domain.models import InspectionFinding, InspectionRecord
from app.modules.maintenance_order.domain.models import WorkOrder
from app.modules.operation_event.domain.models import ProductionEvent, ProductionEventVersion
from app.modules.report.application.workflow import generate_report_map_reduce
from app.modules.report.domain.models import (
    OperationReport,
    ReportExport,
    ReportSource,
    ReportVersion,
)
from app.modules.report.domain.schemas import ReportContentUpdate, ReportCreate
from app.modules.report.infrastructure.repository import ReportRepository
from app.ports.storage import StorageProvider
from app.ports.text import TextLLMProvider
from app.shared.errors import AppError, ConflictError, NotFoundError
from app.shared.platform.models import AIJob
from app.shared.platform.runtime import (
    build_runtime_snapshot,
    system_prompt,
    text_provider_for_job,
)
from app.shared.platform.service import create_job, update_job
from app.shared.security.authorization.dependencies import require_data_scope
from app.shared.security.authorization.permissions import Permissions
from app.shared.security.authorization.schemas import CurrentUser


def digest_scope(scope_filter: dict) -> str:
    raw = json.dumps(scope_filter, sort_keys=True, separators=(",", ":"), ensure_ascii=False)
    return hashlib.sha256(raw.encode()).hexdigest()


async def get_report(session: AsyncSession, report_id: UUID) -> OperationReport:
    report = await ReportRepository(session).get_report(report_id)
    if report is None:
        raise NotFoundError("operation_report", report_id)
    return report


async def create_report(
    session: AsyncSession, payload: ReportCreate, user: CurrentUser
) -> OperationReport:
    require_data_scope(
        user,
        user.organization_unit_id,
        Permissions.REPORT_CREATE,
        owner_id=user.user_id,
    )
    try:
        ZoneInfo(payload.timezone)
    except ZoneInfoNotFoundError as exc:
        raise AppError("INVALID_TIMEZONE", "report timezone is not recognized", 422) from exc
    report = OperationReport(
        report_type=payload.report_type,
        business_date=payload.business_date,
        organization_unit_id=user.organization_unit_id,
        timezone=payload.timezone,
        scope_filter=payload.scope_filter,
        scope_digest=digest_scope(payload.scope_filter),
        status="draft",
        created_by=user.user_id,
    )
    session.add(report)
    try:
        await session.flush()
    except IntegrityError as exc:
        await session.rollback()
        raise ConflictError("report with the same type, date and scope already exists") from exc
    await session.commit()
    return report


async def _collect_sources(session: AsyncSession, report: OperationReport) -> list[dict]:
    org = report.organization_unit_id
    scope = report.scope_filter
    timezone = ZoneInfo(report.timezone)
    local_start = datetime.combine(report.business_date, time.min, tzinfo=timezone)
    utc_start = local_start.astimezone(UTC)
    utc_end = (local_start + timedelta(days=1)).astimezone(UTC)
    sources: list[dict] = []
    audio_query = select(AudioRecord).where(
        AudioRecord.organization_unit_id == org,
        AudioRecord.shift_date == report.business_date,
        AudioRecord.business_status == "confirmed",
        AudioRecord.deleted.is_(False),
    )
    shift_filter = scope.get("shift") or scope.get("shifts")
    if shift_filter and shift_filter != "all":
        shifts = [shift_filter] if isinstance(shift_filter, str) else list(shift_filter)
        audio_query = audio_query.where(AudioRecord.shift_code.in_(shifts))
    audio_rows = await session.scalars(audio_query)
    for audio in audio_rows:
        summary = await session.get(HandoverSummaryVersion, audio.current_summary_version_id)
        if summary:
            sources.append(
                {
                    "id": str(audio.id),
                    "type": "audio_record",
                    "version_id": str(summary.id),
                    "content": summary.content,
                }
            )
    event_query = select(ProductionEvent).where(
        ProductionEvent.organization_unit_id == org,
        ProductionEvent.business_status == "confirmed",
        or_(
            and_(
                ProductionEvent.occurred_at >= utc_start,
                ProductionEvent.occurred_at < utc_end,
            ),
            and_(
                ProductionEvent.occurred_at.is_(None),
                ProductionEvent.created_at >= utc_start,
                ProductionEvent.created_at < utc_end,
            ),
        ),
    )
    event_types = scope.get("event_types")
    if event_types:
        event_query = event_query.where(ProductionEvent.event_type.in_(list(event_types)))
    event_rows = await session.scalars(event_query)
    for event in event_rows:
        version = await session.get(ProductionEventVersion, event.current_version_id)
        sources.append(
            {
                "id": str(event.id),
                "type": "production_event",
                "version_id": str(version.id) if version else None,
                "title": event.title,
                "content": version.description if version else event.title,
            }
        )
    order_rows = await session.scalars(
        select(WorkOrder).where(
            WorkOrder.organization_unit_id == org,
            or_(
                WorkOrder.status.not_in(("closed", "cancelled")),
                and_(WorkOrder.updated_at >= utc_start, WorkOrder.updated_at < utc_end),
            ),
        )
    )
    sources.extend(
        {
            "id": str(order.id),
            "type": "work_order",
            "title": order.title,
            "status": order.status,
            "risk_level": order.risk_level,
        }
        for order in order_rows
    )
    inspection_query = select(InspectionRecord).where(
        InspectionRecord.organization_unit_id == org,
        InspectionRecord.inspected_at >= utc_start,
        InspectionRecord.inspected_at < utc_end,
    )
    if scope.get("station_name"):
        inspection_query = inspection_query.where(
            InspectionRecord.station_name == str(scope["station_name"])
        )
    if scope.get("pipeline_name"):
        inspection_query = inspection_query.where(
            InspectionRecord.pipeline_name == str(scope["pipeline_name"])
        )
    inspections = list(await session.scalars(inspection_query))
    if inspections:
        finding_rows = await session.scalars(
            select(InspectionFinding).where(
                InspectionFinding.inspection_id.in_([item.id for item in inspections]),
                InspectionFinding.review_status == "confirmed",
            )
        )
        sources.extend(
            {
                "id": str(finding.id),
                "type": "inspection_finding",
                "title": finding.title,
                "severity": finding.severity,
                "content": finding.description,
            }
            for finding in finding_rows
        )
    return sources


async def _next_version(session: AsyncSession, report_id: UUID) -> int:
    return (
        await session.scalar(
            select(func.max(ReportVersion.version)).where(ReportVersion.report_id == report_id)
        )
        or 0
    ) + 1


async def generate_report(
    session: AsyncSession,
    *,
    report: OperationReport,
    user: CurrentUser,
    provider: TextLLMProvider,
    inline: bool,
) -> UUID:
    require_data_scope(
        user,
        report.organization_unit_id,
        Permissions.REPORT_GENERATE,
        owner_id=report.created_by,
    )
    job = await create_job(
        session,
        user=user,
        job_type="report_generation",
        resource_type="operation_report",
        resource_id=report.id,
        task_name="app.modules.report.generate_report",
        queue="report",
        config_snapshot=await build_runtime_snapshot(
            session,
            job_type="report_generation",
            provider=provider,
            base={"scope": report.scope_filter},
        ),
        enqueue=not inline,
    )
    if not inline:
        await session.commit()
        return job.id
    await execute_report_generation(
        session,
        job=job,
        report=report,
        user=user,
        provider=provider,
    )
    return job.id


async def execute_report_generation(
    session: AsyncSession,
    *,
    job: AIJob,
    report: OperationReport,
    user: CurrentUser,
    provider: TextLLMProvider,
) -> None:
    if job.status in {"succeeded", "failed", "cancelled"}:
        return
    if job.cancel_requested:
        await update_job(session, job, status="cancelled", progress=job.progress)
        await session.commit()
        return
    await update_job(session, job, status="running", progress=10, message="collecting sources")
    sources = await _collect_sources(session, report)
    provider = text_provider_for_job(session, job, provider)
    prompt = job.config_snapshot.get("prompt") or {}
    draft = await generate_report_map_reduce(
        provider,
        sources,
        report.report_type,
        system_prompt=system_prompt(job.config_snapshot),
        user_template=str(prompt.get("user_template") or "{input}"),
    )
    version = ReportVersion(
        report_id=report.id,
        version=await _next_version(session, report.id),
        source="ai",
        title=draft.title,
        content={"sections": draft.sections, "pending_facts": draft.pending_facts},
        review_status="draft",
        immutable=False,
        created_by=user.user_id,
    )
    session.add(version)
    await session.flush()
    by_id = {str(item.get("id")): item for item in sources}
    for source_id in draft.source_ids:
        item = by_id.get(source_id)
        if item is None:
            continue
        session.add(
            ReportSource(
                report_version_id=version.id,
                source_type=item["type"],
                source_id=UUID(source_id),
                source_version_id=UUID(item["version_id"]) if item.get("version_id") else None,
                snapshot=item,
            )
        )
    report.current_version_id = version.id
    report.status = "draft"
    await update_job(
        session, job, status="succeeded", progress=100, message="report draft persisted"
    )
    await session.commit()


async def execute_report_export(
    session: AsyncSession,
    *,
    job: AIJob,
    report: OperationReport,
    export: ReportExport,
    user: CurrentUser,
    storage: StorageProvider,
) -> None:
    if job.status in {"succeeded", "failed", "cancelled"}:
        return
    if job.cancel_requested:
        export.status = "cancelled"
        await update_job(session, job, status="cancelled", progress=job.progress)
        await session.commit()
        return
    await update_job(session, job, status="running", progress=20, message="rendering DOCX")
    version = await session.get(ReportVersion, export.report_version_id)
    if version is None:
        raise ConflictError("report version is missing")
    document = Document()
    document.add_heading(version.title, 0)
    document.add_paragraph(f"报告日期：{report.business_date}")
    sections = version.content.get("sections", version.content)
    for heading, body in sections.items():
        document.add_heading(str(heading), level=1)
        document.add_paragraph(str(body))
    buffer = io.BytesIO()
    document.save(buffer)
    object_key = f"reports/{report.id}/exports/{export.id}.docx"
    await storage.put_bytes(
        object_key,
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    export.object_key = object_key
    export.status = "succeeded"
    export.completed_at = datetime.now(UTC)
    await update_job(session, job, status="succeeded", progress=100, message="export persisted")
    await session.commit()


async def update_report_content(
    session: AsyncSession,
    *,
    report: OperationReport,
    payload: ReportContentUpdate,
    user: CurrentUser,
) -> ReportVersion:
    require_data_scope(
        user,
        report.organization_unit_id,
        Permissions.REPORT_EDIT,
        owner_id=report.created_by,
    )
    current = await session.get(ReportVersion, report.current_version_id)
    if current and current.immutable:
        raise ConflictError("published report version cannot be overwritten")
    version = ReportVersion(
        report_id=report.id,
        version=await _next_version(session, report.id),
        source="manual",
        title=payload.title,
        content=payload.content,
        review_status="draft",
        immutable=False,
        created_by=user.user_id,
    )
    session.add(version)
    await session.flush()
    report.current_version_id = version.id
    report.status = "draft"
    await session.commit()
    return version


async def restore_report_version(
    session: AsyncSession,
    *,
    report: OperationReport,
    version_no: int,
    user: CurrentUser,
) -> ReportVersion:
    source = await session.scalar(
        select(ReportVersion).where(
            ReportVersion.report_id == report.id, ReportVersion.version == version_no
        )
    )
    if source is None:
        raise NotFoundError("report_version", version_no)
    return await update_report_content(
        session,
        report=report,
        payload=ReportContentUpdate(title=source.title, content=source.content),
        user=user,
    )


async def create_export(
    session: AsyncSession,
    *,
    report: OperationReport,
    export_format: str,
    user: CurrentUser,
    storage: StorageProvider,
    inline: bool,
) -> ReportExport:
    require_data_scope(
        user,
        report.organization_unit_id,
        Permissions.REPORT_EXPORT,
        owner_id=report.created_by,
    )
    if report.current_version_id is None:
        raise ConflictError("report has no version to export")
    if export_format == "pdf":
        raise AppError(
            "PDF_EXPORT_UNAVAILABLE", "PDF export is not enabled in this deployment", 409
        )
    export = ReportExport(
        report_id=report.id,
        report_version_id=report.current_version_id,
        export_format=export_format,
        status="queued",
        requested_by=user.user_id,
    )
    session.add(export)
    await session.flush()
    job = await create_job(
        session,
        user=user,
        job_type="report_export",
        resource_type="report_export",
        resource_id=export.id,
        task_name="app.modules.report.export_report",
        queue="report",
        config_snapshot={"report_id": str(report.id), "format": export_format},
        enqueue=not inline,
    )
    if inline:
        await execute_report_export(
            session,
            job=job,
            report=report,
            export=export,
            user=user,
            storage=storage,
        )
    else:
        await session.commit()
    return export
